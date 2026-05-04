from io import BytesIO
from typing import Any

from fastapi import APIRouter, Depends, HTTPException, Query
from fastapi.responses import Response

from app import database as db
from app.config import settings
from app.deps import get_db

router = APIRouter()


@router.get("/chapters")
async def list_chapters(conn=Depends(get_db)) -> list[dict[str, Any]]:
    try:
        rows = conn.execute(
            "SELECT chapter_id, number, title, arc_goal, status, summary FROM chapters ORDER BY number"
        ).fetchall()
        return [dict(r) for r in rows]
    finally:
        conn.close()


@router.get("/scenes")
async def list_scenes(chapter: int | None = None, conn=Depends(get_db)) -> list[dict[str, Any]]:
    try:
        if chapter:
            rows = conn.execute("""
                SELECT s.scene_id, c.number as chapter_number, s.sequence,
                       s.brief, s.summary, s.status
                FROM scenes s JOIN chapters c ON s.chapter_id = c.chapter_id
                WHERE c.number = ? ORDER BY s.sequence
            """, (chapter,)).fetchall()
        else:
            rows = conn.execute("""
                SELECT s.scene_id, c.number as chapter_number, s.sequence,
                       s.brief, s.summary, s.status
                FROM scenes s JOIN chapters c ON s.chapter_id = c.chapter_id
                ORDER BY c.number, s.sequence
            """).fetchall()
        return [dict(r) for r in rows]
    finally:
        conn.close()


@router.get("/scenes/{scene_id}")
async def get_scene(scene_id: int, conn=Depends(get_db)) -> dict[str, Any]:
    try:
        scene = db.get_scene(conn, scene_id)
        if not scene:
            raise HTTPException(404, f"Scene {scene_id} not found")
        return scene
    finally:
        conn.close()


@router.get("/search")
async def search(q: str = Query(min_length=2), conn=Depends(get_db)) -> list[dict[str, Any]]:
    try:
        return db.search_scenes(conn, q)
    finally:
        conn.close()


# ── Export ────────────────────────────────────────────────────────────────────

@router.get("/export")
async def export_novel(
    format: str = Query(default="docx"),
    conn=Depends(get_db),
) -> Response:
    """Download the complete novel as .docx or .txt.
    Only approved scenes are included; in-progress chapters are skipped."""
    if format not in ("docx", "txt"):
        raise HTTPException(400, "format must be 'docx' or 'txt'")
    try:
        bible = db.get_story_bible(conn)
        title = str(bible.get("title", "Novel"))
        slug = settings.get_active_project()

        chapters = conn.execute(
            "SELECT chapter_id, number, title FROM chapters ORDER BY number"
        ).fetchall()

        chapter_data: list[dict] = []
        for ch in chapters:
            rows = conn.execute(
                "SELECT full_text FROM scenes "
                "WHERE chapter_id=? AND status='approved' ORDER BY sequence",
                (ch["chapter_id"],),
            ).fetchall()
            scenes = [r["full_text"] or "" for r in rows if r["full_text"]]
            if scenes:
                chapter_data.append({
                    "number": ch["number"],
                    "title": ch["title"] or "",
                    "scenes": scenes,
                })

        if not chapter_data:
            raise HTTPException(404, "No approved scenes to export yet.")

        if format == "docx":
            data = _build_docx(title, chapter_data)
            return Response(
                content=data,
                media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                headers={"Content-Disposition": f'attachment; filename="{slug}.docx"'},
            )
        else:
            data = _build_txt(title, chapter_data).encode("utf-8")
            return Response(
                content=data,
                media_type="text/plain; charset=utf-8",
                headers={"Content-Disposition": f'attachment; filename="{slug}.txt"'},
            )
    finally:
        conn.close()


def _build_docx(title: str, chapters: list[dict]) -> bytes:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt

    doc = Document()
    doc.core_properties.title = title

    # Comfortable margins
    for section in doc.sections:
        section.top_margin    = Inches(1.2)
        section.bottom_margin = Inches(1.2)
        section.left_margin   = Inches(1.4)
        section.right_margin  = Inches(1.4)

    # ── Title page ────────────────────────────────────────────────────────────
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.font.size = Pt(28)
    run.font.bold = True
    doc.add_paragraph()   # breathing room before first chapter

    # ── Chapters ─────────────────────────────────────────────────────────────
    for ch in chapters:
        heading_text = f"Kapitel {ch['number']}"
        if ch["title"]:
            heading_text += f" – {ch['title']}"

        h = doc.add_heading(heading_text, level=1)
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER

        for i, scene_text in enumerate(ch["scenes"]):
            if i > 0:
                sep = doc.add_paragraph("*   *   *")
                sep.alignment = WD_ALIGN_PARAGRAPH.CENTER
                sep.paragraph_format.space_before = Pt(12)
                sep.paragraph_format.space_after  = Pt(12)

            first_para_in_scene = True
            for raw in scene_text.split("\n\n"):
                raw = raw.strip()
                if not raw:
                    continue
                p = doc.add_paragraph(raw)
                pf = p.paragraph_format
                pf.space_before = Pt(0)
                pf.space_after  = Pt(0)
                # First paragraph of each scene: no indent; subsequent: first-line indent
                pf.first_line_indent = Pt(0) if first_para_in_scene else Pt(24)
                first_para_in_scene = False

        doc.add_page_break()

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


def _build_txt(title: str, chapters: list[dict]) -> str:
    lines: list[str] = []
    lines.append(title.upper())
    lines.append("=" * max(len(title), 40))
    lines.append("")

    for ch in chapters:
        heading = f"KAPITEL {ch['number']}"
        if ch["title"]:
            heading += f" – {ch['title'].upper()}"
        lines.append("")
        lines.append(heading)
        lines.append("-" * len(heading))
        lines.append("")

        for i, scene_text in enumerate(ch["scenes"]):
            if i > 0:
                lines.append("")
                lines.append("*   *   *")
                lines.append("")
            lines.append(scene_text.strip())

        lines.append("")

    return "\n".join(lines)
